"""
Resume parser service.
Supports PDF (via PyPDF2) and DOCX (via python-docx) formats.
Extracts raw text for further processing by the AI pipeline.
"""
import io
import re
from typing import Optional

import PyPDF2
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from a PDF file.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Extracted text string, or empty string on failure.
    """
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Extracted text string, or empty string on failure.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Parse a resume file and return extracted text.
    Supports .pdf and .docx formats.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename: Original filename (used to determine file type).

    Returns:
        Extracted plain text from the resume.

    Raises:
        ValueError: If the file type is unsupported or parsing fails.
    """
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".doc"):
        raise ValueError("Legacy .doc format is not supported. Please upload a .docx or .pdf file.")
    else:
        raise ValueError(f"Unsupported file type. Please upload a PDF or DOCX file.")

    if not text or len(text.strip()) < 50:
        raise ValueError("Could not extract meaningful text from the resume. The file may be image-based or empty.")

    return clean_text(text)


def clean_text(text: str) -> str:
    """
    Clean extracted resume text:
    - Remove excessive whitespace
    - Normalize line endings
    - Remove non-printable characters
    """
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Remove non-printable characters (but keep newlines and spaces)
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    # Collapse multiple spaces into one
    text = re.sub(r' +', ' ', text)
    # Collapse more than 2 consecutive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_email(text: str) -> Optional[str]:
    """Extract email address from resume text."""
    pattern = r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b'
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number from resume text."""
    pattern = r'[\+]?[\d]?[\s\-\.]?\(?\d{3}\)?[\s\-\.]?\d{3}[\s\-\.]?\d{4}'
    match = re.search(pattern, text)
    return match.group(0).strip() if match else None


def extract_name_heuristic(text: str) -> Optional[str]:
    """
    Heuristic to extract candidate name — looks at the first non-empty lines.
    This is a best-effort extraction; AI parsing will be more accurate.
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:
        # Name likely has 2-4 words, all capitalized, no numbers
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w) and not any(c.isdigit() for c in line):
            return line
    return None
